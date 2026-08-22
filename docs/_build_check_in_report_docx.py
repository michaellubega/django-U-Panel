#!/usr/bin/env python3
"""Build docs/ONLINE_CHECK_IN_REPORT.docx from docs/ONLINE_CHECK_IN_REPORT.md."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

DOCS = Path(__file__).resolve().parent
MD_PATH = DOCS / "ONLINE_CHECK_IN_REPORT.md"
OUT_PATH = DOCS / "ONLINE_CHECK_IN_REPORT.docx"

BRAND = RGBColor(0x17, 0x72, 0x45)


def _ensure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15

    for level, size in ((1, 18), (2, 14), (3, 12)):
        name = f"Heading {level}"
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BRAND
        style.paragraph_format.space_before = Pt(12 if level > 1 else 0)
        style.paragraph_format.space_after = Pt(6)

    if "Code Block" not in [s.name for s in doc.styles]:
        code_style = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Consolas"
        code_style.font.size = Pt(9)
        code_style.paragraph_format.space_before = Pt(4)
        code_style.paragraph_format.space_after = Pt(4)


def _add_runs(paragraph, text: str) -> None:
    """Parse **bold**, `code`, and *italic* inline markers."""
    pattern = re.compile(
        r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*|\[[^\]]+\]\([^)]+\))"
    )
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        chunk = m.group(0)
        if chunk.startswith("**") and chunk.endswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            run.bold = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            run = paragraph.add_run(chunk[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif chunk.startswith("*") and chunk.endswith("*"):
            run = paragraph.add_run(chunk[1:-1])
            run.italic = True
        elif chunk.startswith("["):
            link = re.match(r"\[([^\]]+)\]\(([^)]+)\)", chunk)
            if link:
                run = paragraph.add_run(link.group(1))
                run.font.color.rgb = BRAND
                run.underline = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _parse_table_row(line: str) -> list[str]:
    parts = [c.strip() for c in line.strip().strip("|").split("|")]
    return parts


def _is_table_sep(line: str) -> bool:
    return bool(re.match(r"^\|\s*[-:| ]+\|\s*$", line.strip()))


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(cols):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            text = row[ci] if ci < len(row) else ""
            if ri == 0:
                _add_runs(p, text)
                for run in p.runs:
                    run.bold = True
            else:
                _add_runs(p, text)
    doc.add_paragraph()


def _add_code_block(doc: Document, lines: list[str], lang: str) -> None:
    if lang == "mermaid":
        p = doc.add_paragraph()
        run = p.add_run("Flow diagram (see markdown source for graphical version):")
        run.italic = True
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        for line in lines:
            cp = doc.add_paragraph(style="Code Block")
            cp.add_run(line)
    else:
        for line in lines:
            cp = doc.add_paragraph(style="Code Block")
            cp.add_run(line)


def build_docx(md_text: str, out_path: Path) -> None:
    doc = Document()
    _ensure_styles(doc)

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_lang = ""
    code_lines: list[str] = []
    table_rows: list[list[str]] = []

    def flush_table() -> None:
        nonlocal table_rows
        if table_rows:
            _add_table(doc, table_rows)
            table_rows = []

    while i < len(lines):
        line = lines[i]

        if in_code:
            if line.strip().startswith("```"):
                _add_code_block(doc, code_lines, code_lang)
                code_lines = []
                in_code = False
                code_lang = ""
            else:
                code_lines.append(line)
            i += 1
            continue

        if line.strip().startswith("```"):
            flush_table()
            in_code = True
            code_lang = line.strip()[3:].strip() or "text"
            i += 1
            continue

        if line.strip().startswith("|"):
            if _is_table_sep(line):
                i += 1
                continue
            table_rows.append(_parse_table_row(line))
            i += 1
            continue

        flush_table()

        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, re.sub(r"^\d+\.\s", "", stripped))
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, stripped[2:])
            i += 1
            continue

        p = doc.add_paragraph()
        _add_runs(p, stripped)
        i += 1

    flush_table()
    if in_code and code_lines:
        _add_code_block(doc, code_lines, code_lang)

    doc.save(out_path)
    print(f"Wrote {out_path}")


def main() -> None:
    if not MD_PATH.is_file():
        raise SystemExit(f"Missing {MD_PATH}")
    build_docx(MD_PATH.read_text(encoding="utf-8"), OUT_PATH)


if __name__ == "__main__":
    main()
